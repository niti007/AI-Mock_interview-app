import re
from docx import Document
from pdfminer.high_level import extract_text
import os
from typing import Dict, List, Union


class CVParser:
    def __init__(self):
        self.extracted_text = ""
        self.parsed_data: Dict[str, Union[str, List[str], Dict[str, str]]] = {}

    def parse_cv(self, file_path: str) -> Dict[str, Union[str, List[str], Dict[str, str]]]:
        """
        Parse CV file (PDF or DOCX) from a file path

        Args:
            file_path (str): Path to the CV file

        Returns:
            dict: Parsed CV data

        Raises:
            Exception: If there's an error parsing the file
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            # Extract text based on file type
            file_ext = os.path.splitext(file_path.lower())[1]

            if file_ext == '.pdf':
                self.extracted_text = extract_text(file_path)
            elif file_ext == '.docx':
                self.extracted_text = self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")

            if not self.extracted_text.strip():
                raise ValueError("No text could be extracted from the file")

            # Parse the extracted text
            self.extract_information()
            return self.parsed_data

        except Exception as e:
            raise Exception(f"Error parsing CV: {str(e)}")

    def _extract_docx_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file with improved formatting

        Args:
            file_path (str): Path to the DOCX file

        Returns:
            str: Extracted text with preserved formatting
        """
        doc = Document(file_path)
        full_text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                # Add double newline after headings (usually in all caps)
                if paragraph.text.isupper():
                    full_text.append(f"{paragraph.text}\n")
                else:
                    full_text.append(paragraph.text)

        # Add text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)

        return '\n'.join(full_text)

    def extract_information(self) -> None:
        """Extract relevant information from CV text"""
        self.parsed_data = {
            'skills': self.extract_skills(),
            'education': self.extract_education(),
            'experience': self.extract_experience(),
            'contact': self.extract_contact_info(),
            'additional': self.extract_additional_info()
        }

    def extract_skills(self) -> List[str]:
        """
        Extract skills from CV with improved pattern matching

        Returns:
            list: List of unique skills
        """
        skills_patterns = [
            r"(?i)(SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES|EXPERTISE|QUALIFICATIONS)[:\n](.*?)(?:\n\n|\Z)",
            r"(?i)(TECHNOLOGIES|TOOLS|SOFTWARE)[:\n](.*?)(?:\n\n|\Z)"
        ]

        skills_set = set()
        for pattern in skills_patterns:
            skills_match = re.search(pattern, self.extracted_text, re.DOTALL)
            if skills_match:
                skills_text = skills_match.group(2)
                # Split skills by common separators
                skills = re.split(r'[,•|/\n]', skills_text)
                # Clean and filter empty or whitespace-only skills
                for skill in skills:
                    cleaned_skill = skill.strip()
                    if cleaned_skill and len(cleaned_skill) > 1:  # Avoid single-character skills
                        skills_set.add(cleaned_skill)

        return sorted(list(skills_set))

    def extract_education(self) -> List[Dict[str, str]]:
        """
        Extract education information with structured output

        Returns:
            list: List of education entries with degree, institution, and year
        """
        education_pattern = r"(?i)(EDUCATION|ACADEMIC|QUALIFICATIONS)[:\n](.*?)(?:\n\n|\Z)"
        education_match = re.search(education_pattern, self.extracted_text, re.DOTALL)

        if not education_match:
            return []

        education_text = education_match.group(2)
        education_entries = re.split(r'\n(?=[A-Z])', education_text)

        parsed_education = []
        for entry in education_entries:
            if entry.strip():
                # Try to extract degree, institution and year
                degree_match = re.search(
                    r"(?i)(B\.?S\.?|M\.?S\.?|Ph\.?D\.?|Bachelor'?s?|Master'?s?|Doctorate|MBA|BE|ME|MTech|BTech).*?(?=\n|$)",
                    entry)
                year_match = re.search(r'(19|20)\d{2}(?:\s*-\s*(19|20)\d{2})?', entry)

                parsed_entry = {
                    'degree': degree_match.group(0).strip() if degree_match else '',
                    'institution': re.sub(r'(19|20)\d{2}.*$', '', entry).strip(),
                    'year': year_match.group(0) if year_match else ''
                }
                parsed_education.append(parsed_entry)

        return parsed_education

    def extract_experience(self) -> List[Dict[str, str]]:
        """
        Extract work experience with structured output

        Returns:
            list: List of work experience entries
        """
        experience_pattern = r"(?i)(EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT|PROFESSIONAL EXPERIENCE)[:\n](.*?)(?:\n\n|\Z)"
        experience_match = re.search(experience_pattern, self.extracted_text, re.DOTALL)

        if not experience_match:
            return []

        experience_text = experience_match.group(2)
        experience_entries = re.split(r'\n(?=[A-Z][a-z]+ \d{4}|[A-Z][a-z]+ (19|20)\d{2})', experience_text)

        parsed_experience = []
        for entry in experience_entries:
            if entry.strip():
                # Try to extract company, position, dates and description
                company_match = re.search(r'^([^,\n])+', entry)
                position_match = re.search(r'(?<=,\s)([^,\n])+', entry)
                dates_match = re.search(
                    r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* (?:19|20)\d{2}\s*(?:-|–|to)\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* (?:19|20)\d{2}|Present)',
                    entry)

                description = re.sub(r'^.*\n', '', entry).strip()

                parsed_entry = {
                    'company': company_match.group(0).strip() if company_match else '',
                    'position': position_match.group(0).strip() if position_match else '',
                    'dates': dates_match.group(0) if dates_match else '',
                    'description': description
                }
                parsed_experience.append(parsed_entry)

        return parsed_experience

    def extract_contact_info(self) -> Dict[str, str]:
        """
        Extract contact information with improved pattern matching

        Returns:
            dict: Dictionary containing contact information
        """
        # Extract email with more comprehensive pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email = re.search(email_pattern, self.extracted_text)

        # Extract phone number with international format support
        phone_pattern = r'''(?x)
            (?:
                (?:\+\d{1,3}[-.\s]?)?  # Optional country code
                \(?(?:\d{3})\)?[-.\s]?  # Area code
                \d{3}[-.\s]?\d{4}       # Main number
                |
                \d{4}[-.\s]?\d{3}[-.\s]?\d{3}  # Alternative format
            )
        '''
        phone = re.search(phone_pattern, self.extracted_text, re.VERBOSE)

        # Extract LinkedIn URL
        linkedin_pattern = r'(?:https?:)?\/\/(?:[\w]+\.)?linkedin\.com\/in\/[\w\-\_À-ÿ%]+\/?'
        linkedin = re.search(linkedin_pattern, self.extracted_text)

        # Extract location (city, state/country)
        location_pattern = r'(?i)(?:^|\n)([A-Za-z\s,]+(?:,\s*[A-Za-z\s]+)){1,2}(?=\n|$)'
        location = re.search(location_pattern, self.extracted_text)

        return {
            'email': email.group(0) if email else "",
            'phone': phone.group(0) if phone else "",
            'linkedin': linkedin.group(0) if linkedin else "",
            'location': location.group(0).strip() if location else ""
        }

    def extract_additional_info(self) -> Dict[str, List[str]]:
        """
        Extract additional information such as certifications, languages, etc.

        Returns:
            dict: Dictionary containing additional information
        """
        # Extract certifications
        cert_pattern = r"(?i)(CERTIFICATIONS?|CERTIFICATES?)[:\n](.*?)(?:\n\n|\Z)"
        cert_match = re.search(cert_pattern, self.extracted_text, re.DOTALL)
        certifications = []
        if cert_match:
            cert_text = cert_match.group(2)
            certifications = [cert.strip() for cert in re.split(r'[\n•]', cert_text) if cert.strip()]

        # Extract languages
        lang_pattern = r"(?i)(LANGUAGES?)[:\n](.*?)(?:\n\n|\Z)"
        lang_match = re.search(lang_pattern, self.extracted_text, re.DOTALL)
        languages = []
        if lang_match:
            lang_text = lang_match.group(2)
            languages = [lang.strip() for lang in re.split(r'[,•\n]', lang_text) if lang.strip()]

        return {
            'certifications': certifications,
            'languages': languages
        }